import ast
import os
from docx import Document
from huggingface_hub import login
from PIL.Image import Image
from docx.shared import Inches
from transformers import pipeline

from CodeToFlow_v4 import parse_function, generate_flow_diagram, parse_class


login(token="")

# Load a code explanation model from Hugging Face
code_explainer = pipeline("text-generation", model="bigcode/starcoder")

def generate_explanation(code_snippet):
    prompt = f"Explain the following Python code:\n\n{code_snippet}\n\nExplanation:"
    explanation = code_explainer(prompt, max_length=100, num_return_sequences=1)[0]['generated_text']
    return explanation

def find_entry_point(tree):
    entry_points = []
    for node in ast.walk(tree):
        if isinstance(node, ast.If) and isinstance(node.test, ast.Compare):
            # Check for `if __name__ == "__main__":`
            if (isinstance(node.test.left, ast.Name) and
                node.test.left.id == "__name__" and
                isinstance(node.test.ops[0], ast.Eq) and
                isinstance(node.test.comparators[0], ast.Str) and
                node.test.comparators[0].s == "__main__"):
                entry_points.append(node)
        elif isinstance(node, ast.Expr) and isinstance(node.value, ast.Call):
            # Check for top-level function calls
            entry_points.append(node)
    return entry_points

def create_word_document(functions, classes, code_snippets, tree):
    doc = Document()
    doc.add_heading("Code Flow Diagram and Documentation", level=1)

    # Add entry point information
    entry_points = find_entry_point(tree)
    if entry_points:
        doc.add_heading("Entry Points", level=2)
        for entry in entry_points:
            doc.add_paragraph(ast.unparse(entry))

    # Add diagrams and explanations for functions
    if functions:
        doc.add_heading("Functions", level=2)
        for func in functions:
            # Generate diagram for the function
            nodes = parse_function(func)
            dot = generate_flow_diagram(func.name, nodes)

            # Save the diagram as an image in the images folder
            diagram_image_path = os.path.join("images", f"{func.name}_flow.png")
            dot.render(os.path.join("images", f"{func.name}_flow"), format="png", cleanup=True)

            # Resize the image for better fit in the Word document
            with Image.open(diagram_image_path) as img:
                img = img.resize((800, 600))  # Resize to fit the document
                img.save(diagram_image_path)

            # Add the diagram to the Word document
            doc.add_heading(f"Function: {func.name}", level=3)
            doc.add_picture(diagram_image_path, width=Inches(6))

            # Add explanation and code snippet
            doc.add_paragraph(f"Explanation for function '{func.name}':")
            docstring = ast.get_docstring(func) or generate_explanation(ast.unparse(func))
            doc.add_paragraph(docstring)
            doc.add_paragraph("Code Snippet:")
            try:
                # Use ast.unparse() if available (Python 3.9+)
                doc.add_paragraph(ast.unparse(func))
            except AttributeError:
                # Fallback to ast.dump() for older Python versions
                doc.add_paragraph(ast.dump(func))

    # Add diagrams and explanations for classes
    if classes:
        doc.add_heading("Classes", level=2)
        for cls in classes:
            # Generate diagram for the class
            nodes = parse_class(cls)
            dot = generate_flow_diagram(cls.name, nodes)

            # Save the diagram as an image in the images folder
            diagram_image_path = os.path.join("images", f"{cls.name}_flow.png")
            dot.render(os.path.join("images", f"{cls.name}_flow"), format="png", cleanup=True)

            # Resize the image for better fit in the Word document
            with Image.open(diagram_image_path) as img:
                img = img.resize((800, 800))  # Resize to fit the document
                img.save(diagram_image_path)

            # Add the diagram to the Word document
            doc.add_heading(f"Class: {cls.name}", level=3)
            doc.add_picture(diagram_image_path, width=Inches(6))

            # Add explanation and code snippet
            doc.add_paragraph(f"Explanation for class '{cls.name}':")
            docstring = ast.get_docstring(cls) or generate_explanation(ast.unparse(cls))
            doc.add_paragraph(docstring)
            doc.add_paragraph("Code Snippet:")
            try:
                # Use ast.unparse() if available (Python 3.9+)
                doc.add_paragraph(ast.unparse(cls))
            except AttributeError:
                # Fallback to ast.dump() for older Python versions
                doc.add_paragraph(ast.dump(cls))

    return doc